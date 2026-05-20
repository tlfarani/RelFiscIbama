import streamlit as st
import pandas as pd
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import io
import os
import requests
import zipfile
from datetime import datetime, timedelta

# --- CONFIGURAÇÃO DA PÁGINA ---
st.set_page_config(page_title="IBAMA - Gerador de Relatórios", layout="wide")

# --- CUSTOMIZAÇÃO COMPLETA DE INTERFACE (BLINDAGEM VISUAL) ---
st.markdown("""
    <style>
    /* 1. Títulos e Subtítulos em Verde Musgo */
    h1, h2, h3, .stSubheader, [data-testid="stWidgetLabel"] p {
        color: #4E5D30 !important;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif !important;
        font-weight: bold !important;
    }

    /* 2. Customização das tags internas escolhidas no Multiselect */
    span[data-baseweb="tag"] {
        background-color: #E9EDDE !important;
        color: #4E5D30 !important;
        border: 1px solid #4E5D30 !important;
    }

    /* 3. Estilização Firme dos Botões */
    div.stButton > button:first-child, div.stDownloadButton > button:first-child {
        background-color: #4E5D30 !important;
        color: #FFFFFF !important;
        border-radius: 8px !important;
        border: 1px solid #4E5D30 !important;
        padding: 10px 24px !important;
        font-weight: bold !important;
    }
    div.stButton > button:first-child p, div.stDownloadButton > button:first-child p {
        color: #FFFFFF !important;
    }
    div.stButton > button:first-child:hover, div.stDownloadButton > button:first-child:hover {
        background-color: #3A471E !important;
        border-color: #3A471E !important;
    }

    /* 4. BLINDAGEM DOS FILTROS (Força fundo claro e remove o bloco escuro) */
    div[data-baseweb="select"] > div {
        background-color: #FFFFFF !important;
        color: #000000 !important;
        border: 1px solid #CBD5E1 !important;
    }
    div[data-baseweb="select"] input {
        color: #000000 !important;
    }
    div[data-baseweb="select"] svg {
        fill: #4E5D30 !important;
    }

    /* 5. INJEÇÃO CIRÚRGICA DE CORES ALTERNADAS FIXAS NA TABELA STREAMLIT */
    div[data-testid="stDataEditor"] th {
        background-color: #4E5D30 !important;
        color: #F2F2F2 !important;
        font-weight: bold !important;
    }
    
    div[data-testid="stDataEditor"] tr:nth-child(even) td {
        background-color: #E9EDDE !important;
        color: #000000 !important;
    }
    div[data-testid="stDataEditor"] tr:nth-child(odd) td {
        background-color: #FFFFFF !important;
        color: #000000 !important;
    }
    </style>
""", unsafe_allow_html=True)

# --- FUNÇÕES DE TRATAMENTO ---

def converter_data_excel(valor):
    val_str = str(valor).strip()
    if not val_str or val_str in ["nan", "None", "0"]:
        return " [ DATA - EDITAR MANUAL ] "
    if val_str.isdigit():
        try:
            dias = int(val_str)
            data_real = datetime(1900, 1, 1) + timedelta(days=dias - 2)
            return data_real.strftime("%d/%m/%Y")
        except: pass
    return val_str

def extrair_volume_numerico(valor):
    if pd.isna(valor): return 0.0
    try: return float(str(valor).strip().replace(",", "."))
    except ValueError: return 0.0

def extrair_volume_texto(valor):
    if pd.isna(valor) or str(valor).strip() == "": return " [ VOLUME - EDITAR MANUAL ] "
    val_str = str(valor).strip().lower()
    try:
        num_float = float(val_str.replace(",", "."))
        if num_float == 0.0: return "0"
        texto_formatado = f"{num_float:.7f}"
        if "." in texto_formatado:
            texto_formatado = texto_formatado.rstrip('0').rstrip('.')
        return texto_formatado.replace(".", ",")
    except ValueError: return val_str.replace(".", ",")

# 🌟 FUNÇÃO GLOBAL ÚNICA E TRATADA CONTRA SIEMA FORA DO AR
def t_tag(valor, nome_tag):
    v_s = str(valor).strip()
    if v_s in ["", "nan", "None", "0", "Processo Não Encontrado"]: 
        return f" [ {nome_tag.upper()} - EDITAR MANUAL ] "
    if nome_tag == "siema" and "fora do ar" in v_s.lower():
        return " [ SIEMA FORA DO AR - EDITAR MANUAL ] "
    return v_s

# --- FUNÇÕES DE NEGÓCIO ---

def determinar_jurisdicao(bacia):
    bacia_limpa = str(bacia).lower().strip()
    if "santos" in bacia_limpa: return "-SP"
    if "campos" in bacia_limpa: return "-RJ"
    if "espirito santo" in bacia_limpa: return "-ES"
    return ""

def processar_grandeza(grandeza):
    g = str(grandeza).strip().title()
    tabela = {
        "Potencial": ("quando as consequências não são evidentes", "5"),
        "Reduzida": ("quando os danos ambientais são locais ou temporários", "15"),
        "Fraca": ("quando os danos ambientais são de pequena proporção ou de baixa complexidade, gravidade ou magnitude, diante do contexto considerado", "30"),
        "Moderada": ("quando os danos ambientais são de proporção intermediária ou de moderada complexidade, gravidade ou magnitude, diante do contexto considerado", "50"),
        "Grave": ("quando os danos ambientais são de grande proporção ou de alta complexidade, gravidade ou magnitude, diante do contexto considerado", "70")
    }
    return tabela.get(g, (" [ GRANDEZA TEXTO - EDITAR MANUAL ] ", " [ PONTOS GRANDEZA - EDITAR MANUAL ] "))

def processar_nivel(nivel):
    niveis = {
        "A": "Como o incidente envolveu uma empresa de grande porte, a multa irá variar de, aproximadamente, 150 mil a 10 milhões de reais (Mínimo + 0,3% a 20% do teto)",
        "B": "Como o incidente envolveu uma empresa de grande porte, a multa irá variar de, aproximadamente, 5 milhões a 15 milhões de reais (Mínimo + 10% a 30% do teto)",
        "C": "Como o incidente envolveu uma empresa de grande porte, a multa irá variar de, aproximadamente, 15,5 milhões a 25 milhões de reais (Mínimo + 31% a 50% do teto)",
        "D": "Como o incidente envolveu uma empresa de grande porte, a multa irá variar de, aproximadamente, 25,5 milhões a 37,5 milhões de reais (Mínimo + 51% a 75% do teto)",
        "E": "Como o incidente envolveu uma empresa de grande porte, a multa irá variar de, aproximadamente, 38 milhões a 50 milhões de reais (Mínimo + 76% a 100% do teto)"
    }
    return niveis.get(str(nivel).strip().upper(), " [ NÍVEL TEXTO - EDITAR MANUAL ] ")

def extrair_classe_e_modelo(row):
    class_ol = str(row.get('class_ol', '')).strip().title()
    class_risco_bruto = str(row.get('class_risco', '')).strip().upper()
    vol_num = extrair_volume_numerico(row.get('vol_char', '0'))
    letra_risco = "A"
    for r in ["B", "C", "D", "E"]:
        if r in class_risco_bruto: letra_risco = r; break

    if "Oleoso" in class_ol and "Não" not in class_ol:
        return f"Rel_Fisc_Oleoso_{letra_risco}.docx" if letra_risco in ["A", "B", "D"] else "Rel_Fisc_Oleoso_A.docx", letra_risco
    
    if "Não Oleoso" in class_ol or "Nao Oleoso" in class_ol:
        if letra_risco == "A": return ("Rel_Fisc_Nao_Oleoso_Art_61_A.docx" if vol_num > 8 else "Rel_Fisc_Nao_Oleoso_Art_62_A.docx"), "A"
        if letra_risco == "B": return ("Rel_Fisc_Nao_Oleoso_Art_61_B.docx" if vol_num > 200 else "Rel_Fisc_Nao_Oleoso_Art_62_B.docx"), "B"
        if letra_risco == "D": return "Rel_Fisc_Nao_Oleoso_Art_62_D.docx", "D"
        return "Rel_Fisc_Nao_Oleoso_Art_62_A.docx", "A"
    return None, None

def preencher_documento(caminho_modelo, dicionario_dados):
    doc = Document(caminho_modelo)
    def tratar_p(p):
        for chave, valor in dicionario_dados.items():
            if chave in p.text: p.text = p.text.replace(chave, str(valor))
        if "[" in p.text and "]" in p.text:
            for run in p.runs: run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            
    for p in doc.paragraphs: tratar_p(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs: tratar_p(p)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- INTERFACE ---
st.title("⚖️ Fila de Fiscalização - IBAMA")

@st.cache_data(ttl=300) 
def carregar_dados_sharepoint():
    try:
        url = st.secrets["sharepoint"]["url_planilha"]
        headers = {"Content-Type": "application/json"}
        resposta = requests.post(url, headers=headers, json={})
        resposta.raise_for_status()
        dados_json = resposta.json()
        if isinstance(dados_json, dict) and "value" in dados_json: lista = dados_json["value"]
        elif isinstance(dados_json, list): lista = dados_json
        else: return None
        return pd.DataFrame(lista)
    except Exception as e:
        st.error(f"Erro ao carregar dados: {e}")
        return None

df_original = carregar_dados_sharepoint()

if df_original is not None and not df_original.empty:
    df = df_original.copy()
    
    colunas_map = {
        'ID': 'num_doc', 'PROCESSO': 'processo_sei', 'SIEMA': 'siema', 'SITUACAO': 'situacao',
        'LAUDO_SEI': 'laudo_sei', 'DATA_ACIDENTE': 'data_acid', 'RAIPO_SEI': 'relat_sei',
        'INSTALACAO': 'instalacao', 'Campo': 'campo', 'Bacia': 'bacia', 'EMPRESA': 'empresa',
        'CNPJ': 'cnpj', 'PRODUTO': 'produto', 'CLASS_OL': 'class_ol', 'CLASS_RISCO': 'class_risco',
        'VOL': 'vol_char', 'Lat': 'lat', 'Lon': 'lon', 'Grandeza': 'grandeza',
        'AUTO_INFRACAO': 'auto', 'MULTA_APLICADA': 'multa_char', 'Data_AI': 'data_ai',
        'MULTA_PREVISTA': 'multa_prevista', 'Fiscal': 'fiscal', 'Nivel': 'nivel', 'Nivel_Pontos': 'nivel_pontos',
        'Lat_Auto': 'lat_auto', 'Lon_Auto': 'lon_auto'
    }
    
    for col_real, col_interna in colunas_map.items():
        if col_real in df.columns: df[col_interna] = df[col_real]
        else: df[col_interna] = ""

    # --- FILTROS ---
    with st.container(border=True):
        st.markdown("**🔍 Painel de Filtros**")
        c1, c2, c3 = st.columns(3)
        with c1:
            op_situ = sorted(df['situacao'].astype(str).unique())
            sel_situ = st.multiselect("SITUAÇÃO:", op_situ, default=["Autuar"] if "Autuar" in op_situ else [])
        with c2:
            df['f_limpo'] = df['fiscal'].astype(str).replace({"": "Não Atribuído", "nan": "Não Atribuído", "None": "Não Atribuído"})
            op_fisc = sorted(df['f_limpo'].unique())
            sel_fisc = st.multiselect("FISCAL:", ["Todos"] + op_fisc, default=["Todos"])
        with c3:
            todos_laudos = st.checkbox("Mostrar processos sem LAUDO_SEI", value=False)

    df_f = df.copy()
    if sel_situ: df_f = df_f[df_f['situacao'].astype(str).isin(sel_situ)]
    if sel_fisc and "Todos" not in sel_fisc: df_f = df_f[df_f['f_limpo'].astype(str).isin(sel_fisc)]
    if not todos_laudos: df_f = df_f[df_f['laudo_sei'].astype(str).str.strip() != ""]

    df_f = df_f.reset_index(drop=True)

    # --- CONTROLE DE SELEÇÃO EM MASSA ---
    st.markdown("### 📋 Processos para Análise")
    
    marcar_todos = st.checkbox("✅ Marcar todos os processos mostrados abaixo", value=False)
    vetor_selecao_inicial = [marcar_todos] * len(df_f)

    # --- PREPARAÇÃO DA BASE DE EXIBIÇÃO ---
    df_exib = pd.DataFrame({
        "Selecionar": vetor_selecao_inicial,
        "ID": df_f['num_doc'].astype(str),
        "SITUAÇÃO": df_f['situacao'].astype(str),
        "Fiscal": df_f['f_limpo'].astype(str),
        "Data": [converter_data_excel(d) for d in df_f['data_acid']],
        "Produto": df_f['produto'].astype(str),
        "Class OL": df_f['class_ol'].astype(str),
        "Risco": df_f['class_risco'].astype(str),
        "Vol (m³)": [extrair_volume_texto(v) for v in df_f['vol_char']],
        "Multa Prev": df_f['multa_prevista'].astype(str),
        "Lat Auto": df_f['lat_auto'].astype(str),
        "Lon Auto": df_f['lon_auto'].astype(str)
    })

    tabela_editada = st.data_editor(
        df_exib,
        hide_index=True,
        use_container_width=True,
        disabled=[col for col in df_exib.columns if col != "Selecionar"],
        column_config={
            "Selecionar": st.column_config.CheckboxColumn("Selecionar")
        },
        key=f"editor_{marcar_todos}"
    )
    
    indices_selecionados = tabela_editada[tabela_editada["Selecionar"] == True].index
    selecionados = df_f.iloc[indices_selecionados]

    if not selecionados.empty:
        st.write("---")
        st.subheader(f"🚀 Geração em Lote ({len(selecionados)} itens)")
        
        # --- LÓGICA DE COMPACTAÇÃO EM ZIP (EM MEMÓRIA) ---
        zip_buffer = io.BytesIO()
        arquivos_para_zipar = 0
        
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            for _, row in selecionados.iterrows():
                modelo, risco = extrair_classe_e_modelo(row)
                if not modelo: continue
                
                caminho = os.path.join("modelos", modelo)
                if not os.path.exists(caminho): continue

                grandeza_texto, grandeza_pontos = processar_grandeza(row.get('grandeza', ''))
                nivel_texto = processar_nivel(row.get('nivel', ''))

                dados = {
                    "<<siema>>": t_tag(row.get('siema', ''), "siema"),
                    "<<processo_sei>>": t_tag(row.get('processo_sei', ''), "processo_sei"),
                    "<<laudo_sei>>": str(row.get('laudo_sei', '')).split('.')[0],
                    "<<data_acid>>": converter_data_excel(row.get('data_acid', '')),
                    "<<relat_sei>>": t_tag(row.get('relat_sei', ''), "raipo_sei"),
                    "<<instalacao>>": t_tag(row.get('instalacao', ''), "instalacao"),
                    "<<campo>>": t_tag(row.get('campo', ''), "campo"),
                    "<<bacia>>": t_tag(row.get('bacia', ''), "bacia"),
                    "<<empresa>>": t_tag(row.get('empresa', ''), "empresa"),
                    "<<cnpj>>": t_tag(row.get('cnpj', ''), "cnpj"),
                    "<<produto>>": t_tag(row.get('produto', ''), "produto"),
                    "<<class_ol>>": t_tag(row.get('class_ol', ''), "class_ol"),
                    "<<class_risco>>": risco,
                    "<<vol_char>>": extrair_volume_texto(row.get('vol_char', '')),
                    "<<lat>>": t_tag(row.get('lat', ''), "lat"),
                    "<<lon>>": t_tag(row.get('lon', ''), "lon"),
                    "<<grandeza>>": t_tag(row.get('grandeza', ''), "grandeza"),
                    "<<grandeza_texto>>": grandeza_texto,
                    "<<grandeza_pontos>>": grandeza_pontos,
                    "<<nivel>>": t_tag(row.get('nivel', ''), "nivel"),
                    "<<nivel_pontos>>": t_tag(row.get('nivel_pontos', ''), "nivel_pontos"),
                    "<<nivel_texto>>": nivel_texto,
                    "<<multa_num>>": t_tag(row.get('multa_char', ''), "multa_aplicada"),
                    "<<multa_char>>": t_tag(row.get('multa_char', ''), "multa_aplicada"),
                    "<<data_ai>>": converter_data_excel(row.get('data_ai', '')),
                    "<<auto>>": t_tag(row.get('auto', ''), "auto_infracao"),
                    "<<jurisdicao>>": determinar_jurisdicao(row.get('bacia', ''))
                }
                
                doc_io = preencher_documento(caminho, dados)
                nome_arquivo = f"Rel_Fisc_{row['num_doc']}.docx"
                
                zip_file.writestr(nome_arquivo, doc_io.getvalue())
                arquivos_para_zipar += 1

        zip_buffer.seek(0)
        
        if arquivos_para_zipar > 1:
            st.markdown("### 📦 Download Unificado")
            st.download_button(
                label=f"📥 Baixar Todos os {arquivos_para_zipar} Relatórios (.ZIP)",
                data=zip_buffer,
                file_name=f"relatorios_ibama_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
                mime="application/zip",
                use_container_width=True
            )
            st.write("---")

        st.markdown("### 🔍 Detalhes Individuais dos Itens Selecionados")
        for _, row in selecionados.iterrows():
            modelo, risco = extrair_classe_e_modelo(row)
            if not modelo: continue
            
            caminho = os.path.join("modelos", modelo)
            if not os.path.exists(caminho): continue

            grandeza_texto, grandeza_pontos = processar_grandeza(row.get('grandeza', ''))
            nivel_texto = processar_nivel(row.get('nivel', ''))

            dados_unitarios = {
                "<<siema>>": t_tag(row.get('siema', ''), "siema"),
                "<<processo_sei>>": t_tag(row.get('processo_sei', ''), "processo_sei"),
                "<<laudo_sei>>": str(row.get('laudo_sei', '')).split('.')[0],
                "<<data_acid>>": converter_data_excel(row.get('data_acid', '')),
                "<<relat_sei>>": t_tag(row.get('relat_sei', ''), "raipo_sei"),
                "<<instalacao>>": t_tag(row.get('instalacao', ''), "instalacao"),
                "<<campo>>": t_tag(row.get('campo', ''), "campo"),
                "<<bacia>>": t_tag(row.get('bacia', ''), "bacia"),
                "<<empresa>>": t_tag(row.get('empresa', ''), "empresa"),
                "<<cnpj>>": t_tag(row.get('cnpj', ''), "cnpj"),
                "<<produto>>": t_tag(row.get('produto', ''), "produto"),
                "<<class_ol>>": t_tag(row.get('class_ol', ''), "class_ol"),
                "<<class_risco>>": risco,
                "<<vol_char>>": extrair_volume_texto(row.get('vol_char', '')),
                "<<lat>>": t_tag(row.get('lat', ''), "lat"),
                "<<lon>>": t_tag(row.get('lon', ''), "lon"),
                "<<grandeza>>": t_tag(row.get('grandeza', ''), "grandeza"),
                "<<grandeza_texto>>": grandeza_texto,
                "<<grandeza_pontos>>": grandeza_pontos,
                "<<nivel>>": t_tag(row.get('nivel', ''), "nivel"),
                "<<nivel_pontos>>": t_tag(row.get('nivel_pontos', ''), "nivel_pontos"),
                "<<nivel_texto>>": nivel_texto,
                "<<multa_num>>": t_tag(row.get('multa_char', ''), "multa_aplicada"),
                "<<multa_char>>": t_tag(row.get('multa_char', ''), "multa_aplicada"),
                "<<data_ai>>": converter_data_excel(row.get('data_ai', '')),
                "<<auto>>": t_tag(row.get('auto', ''), "auto_infracao"),
                "<<jurisdicao>>": determinar_jurisdicao(row.get('bacia', ''))
            }
            
            doc_io_unitario = preencher_documento(caminho, dados_unitarios)
            nome = f"Rel_Fisc_{row['num_doc']}.docx"
            
            with st.container(border=True):
                st.write(f"📄 **ID:** {row['num_doc']} | **Processo:** {row['processo_sei']} | **Empresa:** {row['empresa']}")
                st.download_button(label="Baixar Relatório Isolado", data=doc_io_unitario, file_name=nome, key=f"dl_{row['num_doc']}")
else:
    st.info("Aguardando carregamento dos dados do SharePoint...")
